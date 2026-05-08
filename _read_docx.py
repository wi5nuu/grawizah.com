from docx import Document

doc = Document(r"C:\projects\Grawizah\Grawizah_Laporan_Revised_v2 (2).docx")

# Print all paragraphs
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

# Print all tables
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {r_idx}: {' | '.join(cells)}")
