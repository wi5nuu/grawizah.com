import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r"C:\projects\Grawizah\Grawizah_Laporan_Revised_v2 (2).docx")

for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print("\n--- TABLES ---")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx} ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {r_idx}: {' | '.join(cells)}")
